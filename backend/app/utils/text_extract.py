"""Plain-text extraction from uploaded resume files."""
from pathlib import Path

import docx
import pdfplumber


class UnsupportedFileType(Exception):
    pass


def extract_text(file_path: Path, mime_type: str) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".pdf" or "pdf" in mime_type:
        return _extract_pdf(file_path)
    if suffix == ".docx" or "wordprocessingml" in mime_type:
        return _extract_docx(file_path)
    if suffix == ".txt" or mime_type.startswith("text/"):
        return file_path.read_text(encoding="utf-8", errors="ignore")

    raise UnsupportedFileType(f"Unsupported file type: {file_path.name} ({mime_type})")


def _extract_pdf(file_path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            parts.append(text)
    return "\n".join(parts).strip()


def _extract_docx(file_path: Path) -> str:
    document = docx.Document(str(file_path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()
